import os
import io
import tempfile
import mimetypes
from pathlib import Path
from typing import Optional, Dict, List, Tuple
import logging
from datetime import datetime
import threading
import time
import subprocess
import shlex

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from werkzeug.datastructures import FileStorage

from PIL import Image, ImageOps
import pillow_heif

from docx import Document
import PyPDF2
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import mammoth

# Audio support - handle gracefully if not available
AUDIO_SUPPORTED = False
try:
    import pydub
    from pydub import AudioSegment
    AUDIO_SUPPORTED = True
    print("Audio conversion is available")
except ImportError as e:
    print(f"Audio conversion not available: {e}")
    AUDIO_SUPPORTED = False
except Exception as e:
    print(f"Audio conversion setup failed: {e}")
    AUDIO_SUPPORTED = False

app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
MAX_FILE_SIZE = 50 * 1024 * 1024  # Reduced to 50MB for better stability
UPLOAD_FOLDER = '/tmp/uploads'  # Use /tmp for Render
OUTPUT_FOLDER = '/tmp/outputs'  # Use /tmp for Render

# Supported file extensions
ALLOWED_EXTENSIONS = {
    'images': ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff', 'heic', 'heif'],
    'documents': ['pdf', 'doc', 'docx', 'txt', 'html'],
    'video': ['mp4', 'avi', 'mov', 'mkv', 'wmv', 'flv', 'webm']
}

# Add audio extensions only if supported
if AUDIO_SUPPORTED:
    ALLOWED_EXTENSIONS['audio'] = ['mp3', 'wav', 'flac', 'aac', 'ogg', 'm4a']

# Create directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Register HEIF support
try:
    pillow_heif.register_heif_opener()
except:
    print("HEIF support not available")

class FileConverter:
    def __init__(self):
        self.supported_formats = ALLOWED_EXTENSIONS

    def get_file_type(self, filename: str) -> str:
        """Determine file type based on extension"""
        ext = filename.lower().split('.')[-1]
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        return 'unknown'

    def is_supported_file(self, filename: str) -> bool:
        """Check if file type is supported"""
        return self.get_file_type(filename) != 'unknown'

    def convert_file(self, input_path: str, output_path: str, 
                    output_format: str, quality: int = 85, 
                    **kwargs) -> Dict:
        """Main conversion dispatcher"""
        try:
            file_type = self.get_file_type(input_path)
            
            if file_type == 'images':
                return self._convert_image(input_path, output_path, output_format, quality, **kwargs)
            elif file_type == 'documents':
                return self._convert_document(input_path, output_path, output_format, **kwargs)
            elif file_type == 'audio':
                if not AUDIO_SUPPORTED:
                    return {'success': False, 'error': 'Audio conversion not available'}
                return self._convert_audio(input_path, output_path, output_format, quality, **kwargs)
            elif file_type == 'video':
                return self._convert_video(input_path, output_path, output_format, quality, **kwargs)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _convert_image(self, input_path: str, output_path: str, 
                      output_format: str, quality: int = 85, 
                      resize: Optional[Tuple[int, int]] = None,
                      maintain_aspect: bool = True) -> Dict:
        """Convert image files"""
        try:
            with Image.open(input_path) as img:
                # Handle transparency for JPEG
                if output_format.lower() in ['jpg', 'jpeg'] and img.mode in ['RGBA', 'LA']:
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    if img.mode == 'RGBA':
                        background.paste(img, mask=img.split()[-1])
                    else:
                        background.paste(img)
                    img = background
                
                # Resize if requested
                if resize:
                    if maintain_aspect:
                        img.thumbnail(resize, Image.Resampling.LANCZOS)
                    else:
                        img = img.resize(resize, Image.Resampling.LANCZOS)
                
                # Fix orientation
                img = ImageOps.exif_transpose(img)
                
                # Save parameters
                save_params = {'format': output_format.upper()}
                
                if output_format.lower() in ['jpg', 'jpeg']:
                    save_params['quality'] = quality
                    save_params['optimize'] = True
                elif output_format.lower() == 'png':
                    save_params['optimize'] = True
                elif output_format.lower() == 'webp':
                    save_params['quality'] = quality
                    save_params['optimize'] = True
                
                img.save(output_path, **save_params)
                
                output_size = os.path.getsize(output_path)
                
                return {
                    'success': True,
                    'output_size': output_size,
                    'dimensions': img.size,
                    'format': output_format.upper(),
                    'quality': quality
                }
                
        except Exception as e:
            logger.error(f"Image conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _convert_document(self, input_path: str, output_path: str, 
                         output_format: str, **kwargs) -> Dict:
        """Convert document files"""
        try:
            input_ext = input_path.split('.')[-1].lower()
            
            if input_ext == 'docx' and output_format.lower() == 'pdf':
                return self._docx_to_pdf(input_path, output_path)
            elif input_ext == 'docx' and output_format.lower() == 'txt':
                return self._docx_to_txt(input_path, output_path)
            elif input_ext == 'docx' and output_format.lower() == 'html':
                return self._docx_to_html(input_path, output_path)
            elif input_ext == 'txt' and output_format.lower() == 'pdf':
                return self._txt_to_pdf(input_path, output_path)
            elif input_ext == 'pdf' and output_format.lower() == 'txt':
                return self._pdf_to_txt(input_path, output_path)
            else:
                raise ValueError(f"Conversion from {input_ext} to {output_format} not supported")
                
        except Exception as e:
            logger.error(f"Document conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _convert_audio(self, input_path: str, output_path: str, 
                      output_format: str, quality: int = 85, 
                      bitrate: Optional[str] = None) -> Dict:
        """Convert audio files"""
        if not AUDIO_SUPPORTED:
            return {'success': False, 'error': 'Audio conversion not available'}
            
        try:
            audio = AudioSegment.from_file(input_path)
            
            # Set bitrate based on quality
            if not bitrate:
                if quality >= 90:
                    bitrate = "320k"
                elif quality >= 70:
                    bitrate = "192k"
                elif quality >= 50:
                    bitrate = "128k"
                else:
                    bitrate = "96k"
            
            export_params = {
                'format': output_format.lower(),
                'bitrate': bitrate
            }
            
            audio.export(output_path, **export_params)
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'duration': len(audio) / 1000.0,
                'format': output_format.upper(),
                'bitrate': bitrate
            }
            
        except Exception as e:
            logger.error(f"Audio conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _convert_video(self, input_path: str, output_path: str, 
                      output_format: str, quality: int = 85,
                      resolution: Optional[str] = None,
                      fps: Optional[int] = None) -> Dict:
        """Convert video files using FFmpeg"""
        try:
            # Check if FFmpeg is available
            subprocess.run(['ffmpeg', '-version'], capture_output=True, check=True)
            
            cmd = ['ffmpeg', '-i', input_path, '-y']
            
            # Set quality
            if quality >= 90:
                cmd.extend(['-crf', '18'])
            elif quality >= 70:
                cmd.extend(['-crf', '23'])
            elif quality >= 50:
                cmd.extend(['-crf', '28'])
            else:
                cmd.extend(['-crf', '32'])
            
            # Set resolution
            if resolution:
                cmd.extend(['-vf', f'scale={resolution}'])
            
            # Set fps
            if fps:
                cmd.extend(['-r', str(fps)])
            
            # Set codec based on format
            if output_format.lower() == 'mp4':
                cmd.extend(['-codec:v', 'libx264', '-codec:a', 'aac'])
            elif output_format.lower() == 'webm':
                cmd.extend(['-codec:v', 'libvpx-vp9', '-codec:a', 'libopus'])
            elif output_format.lower() == 'avi':
                cmd.extend(['-codec:v', 'libx264', '-codec:a', 'mp3'])
            
            cmd.append(output_path)
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
            
            if result.returncode != 0:
                raise Exception(f"FFmpeg error: {result.stderr}")
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'format': output_format.upper(),
                'quality': quality
            }
            
        except subprocess.CalledProcessError:
            return {'success': False, 'error': 'FFmpeg not available for video conversion'}
        except subprocess.TimeoutExpired:
            return {'success': False, 'error': 'Video conversion timeout'}
        except Exception as e:
            logger.error(f"Video conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _docx_to_pdf(self, input_path: str, output_path: str) -> Dict:
        """Convert DOCX to PDF with better formatting"""
        try:
            # Use ReportLab conversion
            doc = Document(input_path)
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            y = height - 50
            margin = 50
            line_height = 14
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Better text wrapping
                    text = paragraph.text
                    max_chars = int((width - 2 * margin) / 6)  # Approximate chars per line
                    
                    # Split long lines
                    words = text.split()
                    lines = []
                    current_line = ""
                    
                    for word in words:
                        if len(current_line + word + " ") <= max_chars:
                            current_line += word + " "
                        else:
                            if current_line:
                                lines.append(current_line.strip())
                            current_line = word + " "
                    
                    if current_line:
                        lines.append(current_line.strip())
                    
                    # Draw each line
                    for line in lines:
                        c.drawString(margin, y, line)
                        y -= line_height
                        if y < 50:
                            c.showPage()
                            y = height - 50
                    
                    # Add extra space after paragraphs
                    y -= line_height / 2
            
            c.save()
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'format': 'PDF'
            }
            
        except Exception as e:
            logger.error(f"DOCX to PDF conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _docx_to_txt(self, input_path: str, output_path: str) -> Dict:
        """Convert DOCX to TXT"""
        try:
            doc = Document(input_path)
            
            with open(output_path, 'w', encoding='utf-8') as txt_file:
                for paragraph in doc.paragraphs:
                    txt_file.write(paragraph.text + '\n')
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'format': 'TXT'
            }
            
        except Exception as e:
            logger.error(f"DOCX to TXT conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _docx_to_html(self, input_path: str, output_path: str) -> Dict:
        """Convert DOCX to HTML"""
        try:
            with open(input_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                
            with open(output_path, 'w', encoding='utf-8') as html_file:
                html_file.write(result.value)
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'format': 'HTML'
            }
            
        except Exception as e:
            logger.error(f"DOCX to HTML conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _txt_to_pdf(self, input_path: str, output_path: str) -> Dict:
        """Convert TXT to PDF with better formatting"""
        try:
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            y = height - 50
            margin = 50
            line_height = 14
            
            with open(input_path, 'r', encoding='utf-8') as txt_file:
                for line in txt_file:
                    text = line.rstrip('\n\r')
                    if text:
                        # Better text wrapping
                        max_chars = int((width - 2 * margin) / 6)
                        
                        if len(text) <= max_chars:
                            c.drawString(margin, y, text)
                            y -= line_height
                        else:
                            # Split long lines
                            words = text.split()
                            current_line = ""
                            
                            for word in words:
                                if len(current_line + word + " ") <= max_chars:
                                    current_line += word + " "
                                else:
                                    if current_line:
                                        c.drawString(margin, y, current_line.strip())
                                        y -= line_height
                                        if y < 50:
                                            c.showPage()
                                            y = height - 50
                                    current_line = word + " "
                            
                            if current_line:
                                c.drawString(margin, y, current_line.strip())
                                y -= line_height
                    else:
                        # Empty line - add space
                        y -= line_height / 2
                    
                    if y < 50:
                        c.showPage()
                        y = height - 50
            
            c.save()
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'format': 'PDF'
            }
            
        except Exception as e:
            logger.error(f"TXT to PDF conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

    def _pdf_to_txt(self, input_path: str, output_path: str) -> Dict:
        """Convert PDF to TXT"""
        try:
            text = ""
            
            with open(input_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + '\n'
            
            with open(output_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text)
            
            return {
                'success': True,
                'output_size': os.path.getsize(output_path),
                'format': 'TXT'
            }
            
        except Exception as e:
            logger.error(f"PDF to TXT conversion error: {str(e)}")
            return {'success': False, 'error': str(e)}

# Initialize converter
converter = FileConverter()

@app.route('/')
def index():
    """Root endpoint"""
    return jsonify({
        'message': 'File Converter API',
        'version': '1.0.0',
        'endpoints': {
            'convert': '/api/convert',
            'batch_convert': '/api/batch-convert',
            'formats': '/api/formats',
            'health': '/api/health'
        }
    })

@app.route('/api/convert', methods=['POST'])
def convert_file():
    """Convert a single file"""
    try:
        # Check if file is provided
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        # Get parameters
        output_format = request.form.get('format', '').lower()
        quality = int(request.form.get('quality', 85))
        
        resize = request.form.get('resize')
        resolution = request.form.get('resolution')
        fps = request.form.get('fps')
        bitrate = request.form.get('bitrate')
        
        # Validate parameters
        if not output_format:
            return jsonify({'error': 'Output format not specified'}), 400
        
        if not converter.is_supported_file(file.filename):
            return jsonify({'error': 'Unsupported file type'}), 400
        
        # Check file size
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > MAX_FILE_SIZE:
            return jsonify({'error': 'File too large'}), 400
        
        # Save uploaded file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        
        input_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(input_path)
        
        # Create output path
        name_without_ext = os.path.splitext(filename)[0]
        output_filename = f"{name_without_ext}_converted.{output_format}"
        output_path = os.path.join(OUTPUT_FOLDER, f"{timestamp}_{output_filename}")
        
        # Prepare conversion parameters
        conv_params = {'quality': quality}
        
        if resize:
            try:
                width, height = map(int, resize.split('x'))
                conv_params['resize'] = (width, height)
            except ValueError:
                pass
        
        if resolution:
            conv_params['resolution'] = resolution
        
        if fps:
            try:
                conv_params['fps'] = int(fps)
            except ValueError:
                pass
        
        if bitrate:
            conv_params['bitrate'] = bitrate
        
        # Convert file
        result = converter.convert_file(
            input_path, output_path, output_format, **conv_params
        )
        
        # Cleanup input file
        try:
            os.remove(input_path)
        except:
            pass
        
        # Return result
        if result['success']:
            return jsonify({
                'success': True,
                'filename': output_filename,
                'download_url': f'/api/download/{os.path.basename(output_path)}',
                'file_size': result.get('output_size', 0),
                'conversion_info': result
            })
        else:
            return jsonify({'error': result['error']}), 500
            
    except Exception as e:
        logger.error(f"Conversion endpoint error: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/api/download/<filename>')
def download_file(filename):
    """Download converted file"""
    try:
        file_path = os.path.join(OUTPUT_FOLDER, filename)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        mimetype = mimetypes.guess_type(file_path)[0]
        
        return send_file(
            file_path,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        return jsonify({'error': 'Download failed'}), 500

@app.route('/api/batch-convert', methods=['POST'])
def batch_convert():
    """Convert multiple files at once"""
    try:
        files = request.files.getlist('files')
        output_format = request.form.get('format', '').lower()
        quality = int(request.form.get('quality', 85))
        
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        
        if not output_format:
            return jsonify({'error': 'Output format not specified'}), 400
        
        results = []
        
        for file in files:
            if file.filename == '':
                continue
            
            try:
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
                unique_filename = f"{timestamp}_{filename}"
                
                input_path = os.path.join(UPLOAD_FOLDER, unique_filename)
                file.save(input_path)
                
                name_without_ext = os.path.splitext(filename)[0]
                output_filename = f"{name_without_ext}_converted.{output_format}"
                output_path = os.path.join(OUTPUT_FOLDER, f"{timestamp}_{output_filename}")
                
                result = converter.convert_file(
                    input_path, output_path, output_format, quality=quality
                )
                
                # Cleanup input file
                try:
                    os.remove(input_path)
                except:
                    pass
                
                if result['success']:
                    results.append({
                        'original_filename': filename,
                        'converted_filename': output_filename,
                        'download_url': f'/api/download/{os.path.basename(output_path)}',
                        'file_size': result.get('output_size', 0),
                        'success': True
                    })
                else:
                    results.append({
                        'original_filename': filename,
                        'error': result['error'],
                        'success': False
                    })
                    
            except Exception as e:
                results.append({
                    'original_filename': file.filename,
                    'error': str(e),
                    'success': False
                })
        
        return jsonify({
            'results': results,
            'total_files': len(files),
            'successful_conversions': len([r for r in results if r['success']])
        })
        
    except Exception as e:
        logger.error(f"Batch conversion error: {str(e)}")
        return jsonify({'error': 'Batch conversion failed'}), 500

@app.route('/api/formats')
def get_supported_formats():
    """Get supported file formats"""
    return jsonify(ALLOWED_EXTENSIONS)

@app.route('/api/health')
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy', 
        'timestamp': datetime.now().isoformat(),
        'audio_support': AUDIO_SUPPORTED,
        'supported_formats': ALLOWED_EXTENSIONS
    })

def cleanup_old_files():
    """Background task to clean up old files"""
    while True:
        try:
            now = time.time()
            
            # Clean upload folder
            for filename in os.listdir(UPLOAD_FOLDER):
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                if os.path.isfile(file_path):
                    if now - os.path.getmtime(file_path) > 3600:  # 1 hour
                        os.remove(file_path)
                        logger.info(f"Cleaned up old upload: {filename}")
            
            # Clean output folder
            for filename in os.listdir(OUTPUT_FOLDER):
                file_path = os.path.join(OUTPUT_FOLDER, filename)
                if os.path.isfile(file_path):
                    if now - os.path.getmtime(file_path) > 3600:  # 1 hour
                        os.remove(file_path)
                        logger.info(f"Cleaned up old output: {filename}")
                        
        except Exception as e:
            logger.error(f"Cleanup error: {str(e)}")
        
        # Sleep for 5 minutes
        time.sleep(300)

# Start cleanup thread
cleanup_thread = threading.Thread(target=cleanup_old_files, daemon=True)
cleanup_thread.start()

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    logger.info("Starting File Converter API")
    logger.info(f"Audio support: {AUDIO_SUPPORTED}")
    logger.info(f"Supported formats: {ALLOWED_EXTENSIONS}")
    
    app.run(host='0.0.0.0', port=port)